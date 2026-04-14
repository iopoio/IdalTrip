import docx
import os

manual_path = r"c:\Projects\Automation\IdalTrip\download\개방데이터_활용매뉴얼(국문)\한국관광공사_개방데이터_활용매뉴얼(국문)_v4.4.docx"
target_key = "9fae9e07e42ee40cb37ad727972a65bad4fd9117b52a17497004bc89cf12d3c5"

doc = docx.Document(manual_path)
found = False

print(f"Searching for key in manual: {target_key[:10]}...")

for i, para in enumerate(doc.paragraphs):
    if target_key in para.text:
        print(f"!!! FOUND KEY in paragraph {i} !!!")
        print(f"Context: {para.text}")
        found = True
        break

if not found:
    print("The key does not exist in the manual paragraphs.")
    # Check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if target_key in cell.text:
                    print(f"!!! FOUND KEY in a TABLE !!!")
                    found = True
                    break

if not found:
    print("The key is NOT in the manual at all.")
